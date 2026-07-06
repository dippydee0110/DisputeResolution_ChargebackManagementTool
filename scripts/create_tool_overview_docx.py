from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches

BASE_DIR = Path(r"c:\Users\suman\Documents\Dispute & Chargeback Management Agent")
OUTPUT_DIR = BASE_DIR / "output"
SCREENSHOT_DIR = OUTPUT_DIR / "screenshots"
OUT_PATH = OUTPUT_DIR / "TOOL_PROCESS_OVERVIEW.docx"


def add_heading_and_text(doc: Document, heading: str, lines: list[str], level: int = 1) -> None:
    doc.add_heading(heading, level=level)
    for line in lines:
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def add_image_section(doc: Document, heading: str, image_path: Path, caption: str) -> None:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(caption)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.8))
    else:
        doc.add_paragraph(f"Image not found: {image_path.name}")


def main() -> None:
    doc = Document()

    doc.add_heading("Early Dispute Detection and Chargeback Management", level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(
        "This document explains the end-to-end operational process, decisioning workflow, "
        "system components, and visual dashboard context for dispute and chargeback handling."
    )

    add_heading_and_text(
        doc,
        "1. Process Objectives",
        [
            "- Detect dispute signals before formal chargebacks are filed.",
            "- Prioritize response using risk tiers (Critical, High, Medium, Low).",
            "- Coordinate customer, merchant, and admin workflows with clear handoffs.",
            "- Keep audit-ready records for status transitions, evidence, and outcomes.",
        ],
    )

    add_heading_and_text(
        doc,
        "2. End-to-End Operational Process",
        [
            "1) Evidence Intake",
            "- Capture chat transcripts in evidence/CASEID/chat_CASEID.txt.",
            "",
            "2) Signal Detection and Scoring",
            "- Analyze conversation text for dispute indicators.",
            "- Compute weighted risk score and classify risk tier.",
            "",
            "3) Triage and Assignment",
            "- Route cases to customer and merchant resolution workflows.",
            "- Prioritize high-risk cases for immediate action.",
            "",
            "4) Negotiation and Evidence Handling",
            "- Track notes, amount context, evidence requests, and response outcomes.",
            "",
            "5) Escalation and Chargeback Handling",
            "- Escalate unresolved or rejected disputes into chargeback processing.",
            "",
            "6) Closure and Reporting",
            "- Record final outcomes, update dashboards, and publish reporting artifacts.",
        ],
    )

    add_heading_and_text(
        doc,
        "3. Application Layers and Functions",
        [
            "- Data Layer: generate_chargeback_dataset.py creates normalized chargeback records.",
            "- Detection Layer: early_dispute_detection.py extracts signals and computes risk.",
            "- Orchestration Layer: run_analysis.py runs complete generation flow.",
            "- Reporting Layer: HTML dashboards, JSON output, PDF/PPT deliverables.",
            "- Governance Layer: merchant admin controls, intervention and escalation actions.",
        ],
    )

    add_heading_and_text(
        doc,
        "4. Decisioning Workflow (Agent Loop)",
        [
            "- Input event enters the agent orchestrator with case context.",
            "- LLM planner decides whether tool calls are needed.",
            "- Tools execute for signal extraction, evidence checks, policy checks, and scoring.",
            "- LLM reasoner composes recommendation and rationale.",
            "- Confidence and policy gate decides auto-action vs human review.",
            "- Final decision is executed, logged, and looped back until case closure.",
        ],
    )

    add_heading_and_text(
        doc,
        "5. Outputs Produced",
        [
            "- output/early_dispute_detection.html: case-level risk and recommendations.",
            "- output/dispute_detection_results.json: machine-readable analysis output.",
            "- output/chargeback_records.xlsx: operations dataset.",
            "- output/TOOL_OVERVIEW.pdf: visual process summary.",
            "- output/TOOL_OVERVIEW_PRESENTATION.pptx: stakeholder presentation.",
        ],
    )

    doc.add_heading("6. Screenshot and Diagram Walkthrough", level=1)

    add_image_section(
        doc,
        "6.1 Application Context: Entry and Navigation",
        SCREENSHOT_DIR / "index.png",
        "System portal showing role-based entry points and workflow access.",
    )

    add_image_section(
        doc,
        "6.2 Merchant Admin Dashboard",
        SCREENSHOT_DIR / "merchant_dashboard.png",
        "Merchant-focused operations view for case workflow, evidence readiness, and issuer simulation.",
    )

    add_image_section(
        doc,
        "6.3 Early Dispute Detection Dashboard",
        SCREENSHOT_DIR / "early_dispute_detection.png",
        "Risk-ranked detection view with signal traceability and triage guidance.",
    )

    add_image_section(
        doc,
        "6.4 Resolution Center: Customer",
        SCREENSHOT_DIR / "resolution_center_customer.png",
        "Customer-side actions including notes, evidence, amount context, and decision response.",
    )

    add_image_section(
        doc,
        "6.5 Resolution Center: Merchant",
        SCREENSHOT_DIR / "resolution_center_merchant.png",
        "Merchant-side representment and settlement progression with workflow state tracking.",
    )

    add_image_section(
        doc,
        "6.6 Workflow Diagram: Agentic Detection Flow",
        OUTPUT_DIR / "agentic_workflow_diagram.png",
        "Detection flow from chat input through weighted scoring, risk classification, and outputs.",
    )

    add_image_section(
        doc,
        "6.7 Workflow Diagram: Agent Loop (LLM + Tools)",
        OUTPUT_DIR / "agent_loop_llm_tools_workflow.png",
        "Agent-loop decisioning with tool routing, confidence gates, human fallback, and state persistence.",
    )

    add_heading_and_text(
        doc,
        "7. Recommended Operating Cadence",
        [
            "- Daily: run analysis and review high/critical cases first.",
            "- Weekly: review trend shifts and false-positive patterns.",
            "- Monthly: tune thresholds, policy mappings, and intervention rules.",
        ],
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
